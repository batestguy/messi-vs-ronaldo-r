#!/usr/bin/env python3
"""Validate matching retained Word and PDF explanation deliverables."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from pypdf import PdfReader


PLACEHOLDERS = ("TODO", "TBD", "LOREM IPSUM", "PLACEHOLDER")


def normalized(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def docx_text(path: Path) -> tuple[str, Document]:
    document = Document(path)
    pieces = [paragraph.text for paragraph in document.paragraphs]
    pieces.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return normalized("\n".join(pieces)), document


def pdf_text(path: Path) -> tuple[str, PdfReader]:
    reader = PdfReader(str(path))
    pieces = [(page.extract_text() or "") for page in reader.pages]
    return normalized("\n".join(pieces)), reader


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Check paired .docx/.pdf explanation files."
    )
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--pdf", required=True, type=Path)
    parser.add_argument("--required", action="append", default=[])
    parser.add_argument("--min-pages", type=int, default=1)
    parser.add_argument(
        "--allow-different-stems",
        action="store_true",
        help="Permit Word and PDF filenames with different stems.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    errors: list[str] = []

    for path, suffix in ((args.docx, ".docx"), (args.pdf, ".pdf")):
        if not path.is_file():
            errors.append(f"Missing file: {path}")
        elif path.suffix.lower() != suffix:
            errors.append(f"Expected {suffix} file: {path}")
        elif path.stat().st_size == 0:
            errors.append(f"Empty file: {path}")

    if errors:
        print("\n".join(errors), file=sys.stderr)
        return 1

    if not args.allow_different_stems and args.docx.stem != args.pdf.stem:
        errors.append("Word and PDF filenames must use the same stem.")

    try:
        word_text, document = docx_text(args.docx)
    except Exception as exc:
        errors.append(f"Could not open Word file: {exc}")
        word_text, document = "", None

    try:
        portable_text, reader = pdf_text(args.pdf)
    except Exception as exc:
        errors.append(f"Could not open PDF file: {exc}")
        portable_text, reader = "", None

    if document is not None:
        headings = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style is not None
            and paragraph.style.name.startswith("Heading")
        ]
        if not word_text:
            errors.append("Word file has no extractable text.")
        if not headings:
            errors.append("Word file has no real Heading styles.")

    if reader is not None:
        if len(reader.pages) < args.min_pages:
            errors.append(
                f"PDF has {len(reader.pages)} page(s); expected at least "
                f"{args.min_pages}."
            )
        if not portable_text:
            errors.append("PDF has no extractable text.")
        for index, page in enumerate(reader.pages, start=1):
            box = page.mediabox
            if float(box.width) <= 0 or float(box.height) <= 0:
                errors.append(f"PDF page {index} has invalid dimensions.")

    upper_word = word_text.upper()
    upper_pdf = portable_text.upper()
    for token in PLACEHOLDERS:
        if token in upper_word or token in upper_pdf:
            errors.append(f"Placeholder token found: {token}")

    for phrase in args.required:
        if normalized(phrase) not in word_text:
            errors.append(f"Required phrase missing from Word: {phrase}")
        if normalized(phrase) not in portable_text:
            errors.append(f"Required phrase missing from PDF: {phrase}")

    if errors:
        print("\n".join(f"ERROR: {error}" for error in errors), file=sys.stderr)
        return 1

    print(
        "Validated matching explanation files: "
        f"Word={args.docx.stat().st_size:,} bytes; "
        f"PDF={args.pdf.stat().st_size:,} bytes; "
        f"PDF pages={len(reader.pages)}; "
        f"Word headings={len(headings)}."
    )
    print("Text/package checks passed; visual page inspection is still required.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
