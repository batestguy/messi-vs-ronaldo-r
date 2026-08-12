#!/usr/bin/env python3
"""Build the Wave 5 statistics teaching guide as a verified Word intermediate.

Run from any working directory:

    python scripts/build_statistics_guide.py

The script reads the committed analysis artifacts, asks plain Rscript to extract
facts from the shipped analysis artifacts, asserts the documented Wave 5
invariants, and writes the Word intermediate used for the retained PDF:

    output/word/Messi_vs_Ronaldo_Statistics_Explained.docx

It does not modify analysis data, app code, or dashboard behavior.
"""

from __future__ import annotations

import csv
import math
import os
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "word" / "Messi_vs_Ronaldo_Statistics_Explained.docx"

NAVY = "15324B"
MESSI = "14447D"
RONALDO = "A61E2C"
TEAL = "1C7C7D"
GOLD = "B98216"
INK = "263746"
MUTED = "586A78"
PALE_BLUE = "EAF2FA"
PALE_RED = "FBEDEF"
PALE_GOLD = "FFF5D9"
PALE_TEAL = "E9F5F4"
PALE_GRAY = "F2F5F7"
WHITE = "FFFFFF"
RULE = "CAD5DE"
TOTAL_PAGES = 29


R_FACT_EXPORTER = r'''
suppressPackageStartupMessages(library(data.table))

P_ORDER <- c("Messi", "Ronaldo")
source("app/R/mod_weighting.R")
source("app/R/mod_inference.R")
source("app/R/mod_summary.R")

out <- Sys.getenv("GUIDE_FACTS_PATH")
if (!nzchar(out)) stop("GUIDE_FACTS_PATH is not set")

b <- readRDS("data/processed/analysis_bundle.rds")
f <- readRDS("data/processed/famd_loadings.rds")
raw_goal <- readRDS("data/raw/fbref_goal_logs.rds")
raw_match <- readRDS("data/raw/fbref_match_logs.rds")
raw_season <- readRDS("data/raw/fbref_season_stats.rds")
raw_understat <- readRDS("data/raw/understat_goals.rds")
raw_transfer <- readRDS("data/raw/transfermarkt_transfers.rds")
raw_bio <- readRDS("data/raw/transfermarkt_bio.rds")
master_header <- names(fread("data/processed/goals_master_final.csv", nrows = 0))

facts <- list()
add <- function(key, value) facts[[key]] <<- paste(value, collapse = "||")

add("raw_goal_rows", nrow(raw_goal))
add("raw_match_rows", nrow(raw_match))
add("raw_season_rows", nrow(raw_season))
add("raw_understat_rows", nrow(raw_understat))
add("raw_transfer_rows", nrow(raw_transfer))
add("raw_bio_rows", nrow(raw_bio))
add("club_elo_rows", nrow(fread("data/processed/club_elo_lookup.csv", showProgress = FALSE)))
add("national_elo_rows", nrow(fread("data/processed/national_elo_lookup.csv", showProgress = FALSE)))

raw_minutes <- suppressWarnings(as.numeric(raw_match$Min))
header_junk <- sum(raw_match$Opponent == "Opponent" | raw_match$Gls == "Gls", na.rm = TRUE)
pre_dedup <- raw_match[!is.na(raw_minutes) & raw_match$Opponent != "Opponent", ]
pre_opp <- sub("^[a-z]{2,3} ", "", pre_dedup$Opponent)
match_key <- paste(pre_dedup$Player, as.Date(pre_dedup$Date), pre_opp,
                   pre_dedup$Comp, pre_dedup$Round)
duplicates <- sum(duplicated(match_key))
dnp <- nrow(raw_match) - header_junk - nrow(pre_dedup)
add("header_junk", header_junk)
add("dnp_rows", dnp)
add("duplicate_rows", duplicates)
add("pre_dedup_rows", nrow(pre_dedup))

goals <- b$goals
matches <- b$valid_matches
add("goal_rows", nrow(goals))
add("valid_match_rows", nrow(matches))
add("total_minutes", sum(matches$Minutes))
add("scoreless_total", sum(matches$Gls == 0))
add("xg_matched", sum(!is.na(goals$xg)))
add("xg_missing", sum(is.na(goals$xg)))
add("difficulty_present", sum(!is.na(goals$Difficulty_Score)))
add("difficulty_missing", sum(is.na(goals$Difficulty_Score)))
add("master_fields", master_header)
add("master_field_count", length(master_header))
derived <- setdiff(names(goals), master_header)
add("derived_fields", derived)
add("derived_field_count", length(derived))
add("match_fields", names(matches))
add("match_field_count", ncol(matches))

for (p in c("Messi", "Ronaldo")) {
  pk <- tolower(p)
  pg <- goals[Player == p]
  pm <- matches[Player == p]
  add(paste0(pk, "_goals"), nrow(pg))
  add(paste0(pk, "_apps"), nrow(pm))
  add(paste0(pk, "_minutes"), sum(pm$Minutes))
  add(paste0(pk, "_scoreless"), sum(pm$Gls == 0))
  add(paste0(pk, "_penalties"), sum(pg$Penalty))
  add(paste0(pk, "_difficulty_missing"), sum(is.na(pg$Difficulty_Score)))
  add(paste0(pk, "_raw90"), 90 * nrow(pg) / sum(pm$Minutes))
  add(paste0(pk, "_raw90_nopen"), 90 * sum(!pg$Penalty) / sum(pm$Minutes))
  add(paste0(pk, "_weighted_sum"), sum(pg$Difficulty_Score, na.rm = TRUE))
  add(paste0(pk, "_weighted90"), 90 * sum(pg$Difficulty_Score, na.rm = TRUE) / sum(pm$Minutes))
  add(paste0(pk, "_weighted90_nopen"), 90 * sum(pg$Difficulty_Score[!pg$Penalty], na.rm = TRUE) / sum(pm$Minutes))
}

for (src in c("club", "national", "league_avg", "global_1500")) {
  add(paste0("match_elo_", src), sum(matches$Elo_Source == src))
}
add("match_elo_direct", sum(matches$Elo_Source %in% c("club", "national")))
add("match_elo_direct_or_league", sum(matches$Elo_Source %in% c("club", "national", "league_avg")))

add("famd_dim1_variance", b$famd_info$eig[1, 2])
for (nm in rownames(f$var$contrib)) {
  key <- gsub("[^A-Za-z0-9]+", "_", tolower(nm))
  add(paste0("famd_contrib_", key), f$var$contrib[nm, 1])
}

scored <- goals[!is.na(Difficulty_Score)]
scored[, Difficulty_Rank := rank(Difficulty_Score, ties.method = "first")]
scored[, Decile_Number := pmin(10L, floor((Difficulty_Rank - 1) * 10 / .N) + 1L)]
add("d10_min", min(scored[Decile_Number == 10, Difficulty_Score]))
add("d10_max", max(scored[Decile_Number == 10, Difficulty_Score]))

mins <- matches[, .(Minutes = sum(Minutes)), by = Player]
for (excluded in c(FALSE, TRUE)) {
  eligible <- scored
  suffix <- if (excluded) "nopen" else "all"
  if (excluded) eligible <- eligible[Penalty == FALSE]
  for (p in c("Messi", "Ronaldo")) {
    pk <- tolower(p)
    ps <- eligible[Player == p]
    add(paste0(pk, "_eligible_", suffix), nrow(ps))
    add(paste0(pk, "_median_", suffix), median(ps$Difficulty_Score))
    add(paste0(pk, "_d10_", suffix), sum(ps$Decile_Number == 10))
    add(paste0(pk, "_d10_share_", suffix), mean(ps$Decile_Number == 10))
  }
  gaps <- numeric()
  for (k in seq(0.5, 3.0, 0.1)) {
    vals <- sapply(c("Messi", "Ronaldo"), function(p) {
      s <- eligible[Player == p, Difficulty_Score]
      denominator <- mins[Player == p, Minutes]
      90 * sum(sign(s) * abs(s)^k) / denominator
    })
    gaps <- c(gaps, vals[1] - vals[2])
    if (k %in% c(0.5, 1.0, 1.5, 3.0)) {
      kkey <- gsub("\\.", "_", sprintf("%.1f", k))
      add(paste0("k_", kkey, "_messi_", suffix), vals[1])
      add(paste0("k_", kkey, "_ronaldo_", suffix), vals[2])
      add(paste0("k_", kkey, "_gap_", suffix), vals[1] - vals[2])
    }
  }
  add(paste0("stability_", suffix), all(gaps > 0))
  add(paste0("gap_min_", suffix), min(gaps))
  add(paste0("gap_max_", suffix), max(gaps))
}

club_targets <- list(
  "club_messi_barcelona" = c("Messi", "Barcelona"),
  "club_messi_psg" = c("Messi", "PSG"),
  "club_messi_inter_miami" = c("Messi", "Inter Miami"),
  "club_ronaldo_sporting" = c("Ronaldo", "Sporting CP"),
  "club_ronaldo_man_utd" = c("Ronaldo", "Manchester Utd"),
  "club_ronaldo_real_madrid" = c("Ronaldo", "Real Madrid"),
  "club_ronaldo_juventus" = c("Ronaldo", "Juventus"),
  "club_ronaldo_al_nassr" = c("Ronaldo", "Al-Nassr")
)
for (nm in names(club_targets)) {
  v <- club_targets[[nm]]
  add(nm, nrow(goals[Player == v[1] & Squad_clean == v[2]]))
}

h2h_keys <- c("Player", "Date", "Comp", "Round", "Opp_clean")
add("competition_count", uniqueN(matches$Comp))
add("venue_count", uniqueN(matches$Venue))
add("h2h_match_key_duplicates", matches[, .N, by = h2h_keys][N > 1, .N])

goal_keys <- unique(goals[, ..h2h_keys])
match_keys <- unique(matches[, ..h2h_keys])
unmatched_goal_keys <- fsetdiff(goal_keys, match_keys)
add("h2h_unmatched_goal_keys", nrow(unmatched_goal_keys))

contributions <- goals[!is.na(Difficulty_Score), .(
  Weighted = sum(Difficulty_Score),
  Eligible_Goals = .N
), by = h2h_keys]
trajectory <- merge(
  matches,
  contributions,
  by = h2h_keys,
  all.x = TRUE,
  sort = FALSE
)
trajectory[is.na(Weighted), `:=`(Weighted = 0, Eligible_Goals = 0L)]
setorder(trajectory, Player, Date, Comp, Round, Opp_clean)
trajectory[, `:=`(
  Cumulative = cumsum(Weighted),
  Rolling_Weighted = frollsum(Weighted, 30L, align = "right", fill = NA_real_),
  Rolling_Minutes = frollsum(Minutes, 30L, align = "right", fill = NA_real_)
), by = Player]
trajectory[, Rolling30 := 90 * Rolling_Weighted / Rolling_Minutes]

for (p in c("Messi", "Ronaldo")) {
  pk <- tolower(p)
  pg <- goals[Player == p]
  pt <- trajectory[Player == p]
  add(paste0(pk, "_h2h_cumulative"), tail(pt$Cumulative, 1))
  add(paste0(pk, "_h2h_rolling30"), tail(pt$Rolling30[is.finite(pt$Rolling30)], 1))
  add(paste0(pk, "_h2h_scatter"), nrow(pg[is.finite(Difficulty_Score) & is.finite(Opponent_Elo)]))
  add(paste0(pk, "_h2h_distinct_elo"), uniqueN(pg[is.finite(Difficulty_Score) & is.finite(Opponent_Elo), Opponent_Elo]))
  add(paste0(pk, "_open_play"), sum(!pg$Penalty))
  add(paste0(pk, "_penalty_share"), mean(pg$Penalty))
}

scope_specs <- list(
  champions_away = c("Champions Lg", "Away"),
  world_cup_neutral = c("World Cup", "Neutral"),
  mls_all = c("MLS", "All venues"),
  fa_cup_all = c("FA Cup", "All venues"),
  mls_neutral = c("MLS", "Neutral")
)
for (scope_name in names(scope_specs)) {
  spec <- scope_specs[[scope_name]]
  sm <- matches[Comp == spec[1]]
  sg <- goals[Comp == spec[1]]
  if (spec[2] != "All venues") {
    sm <- sm[Venue == spec[2]]
    sg <- sg[Venue == spec[2]]
  }
  for (p in c("Messi", "Ronaldo")) {
    pk <- tolower(p)
    add(paste0(scope_name, "_", pk, "_apps"), nrow(sm[Player == p]))
    add(paste0(scope_name, "_", pk, "_goals"), nrow(sg[Player == p]))
  }
}

wave4 <- inference_run_analysis(
  goals,
  matches,
  inference_snapshot(1, FALSE, "All competitions", "All venues")
)
add("wave4_gap", wave4$overall$observed_gap)
add("wave4_ci_low", wave4$overall$interval[1])
add("wave4_ci_high", wave4$overall$interval[2])
add("wave4_probability", wave4$overall$probability_above_zero)
add("wave4_cohens_d", wave4$overall$cohens_d)
add("wave4_reps", length(wave4$overall$differences))
add("wave4_finite", sum(is.finite(wave4$overall$differences)))
add("wave4_seed", INFERENCE_SEED)

add_subgroup <- function(prefix, tab) {
  for (i in seq_len(nrow(tab))) {
    slug <- gsub("_+$", "", gsub("[^a-z0-9]+", "_", tolower(tab$Subgroup[i])))
    add(paste(prefix, slug, "messi_apps", sep = "_"), tab$Messi_Appearances[i])
    add(paste(prefix, slug, "messi_rate", sep = "_"), tab$Messi_Rate[i])
    add(paste(prefix, slug, "ronaldo_apps", sep = "_"), tab$Ronaldo_Appearances[i])
    add(paste(prefix, slug, "ronaldo_rate", sep = "_"), tab$Ronaldo_Rate[i])
    add(paste(prefix, slug, "gap", sep = "_"), tab$Gap[i])
    add(paste(prefix, slug, "ci_low", sep = "_"), tab$CI_Low[i])
    add(paste(prefix, slug, "ci_high", sep = "_"), tab$CI_High[i])
    add(paste(prefix, slug, "cohens_d", sep = "_"), tab$Cohens_d[i])
    add(paste(prefix, slug, "sparse", sep = "_"), tab$Sparse[i])
  }
}
add_subgroup("era", wave4$era)
add_subgroup("family", wave4$family)

wave5_summary <- summary_comparison_table(
  summary_prepare_scope(goals, matches, summary_snapshot())
)
add("wave5_summary_rows", nrow(wave5_summary))
add_summary_row <- function(prefix, row) {
  add(paste0(prefix, "_messi_apps"), row$Apps_Messi)
  add(paste0(prefix, "_messi_minutes"), row$Minutes_Messi)
  add(paste0(prefix, "_ronaldo_apps"), row$Apps_Ronaldo)
  add(paste0(prefix, "_ronaldo_minutes"), row$Minutes_Ronaldo)
  add(paste0(prefix, "_messi_weighted90"), row$Weighted_per90_Messi)
  add(paste0(prefix, "_ronaldo_weighted90"), row$Weighted_per90_Ronaldo)
  add(paste0(prefix, "_weighted_gap"), row$Weighted_Gap)
  add(paste0(prefix, "_messi_raw90"), row$Raw_per90_Messi)
  add(paste0(prefix, "_ronaldo_raw90"), row$Raw_per90_Ronaldo)
  add(paste0(prefix, "_raw_gap"), row$Raw_Gap)
}
add_summary_row("summary_overall", wave5_summary[Section == "Overall"])
for (venue in SUMMARY_VENUE_LEVELS) {
  prefix <- paste0("summary_venue_", tolower(venue))
  add_summary_row(prefix, wave5_summary[Section == "Venue" & Summary_Scope == venue])
}

raw_csv <- "data/processed/goals_master_final.csv"
app_csv <- "app/data/goals_master_final.csv"
add("raw_csv_rows", nrow(fread(raw_csv, colClasses = "character", na.strings = NULL)))
add("raw_csv_columns", length(master_header))
add("raw_csv_bytes", file.info(raw_csv)$size)
add("raw_csv_md5", unname(tools::md5sum(raw_csv)))
add("app_csv_md5", unname(tools::md5sum(app_csv)))

stopifnot(
  nrow(raw_goal) == 1738L,
  nrow(raw_match) == 2263L,
  header_junk == 12L,
  dnp == 47L,
  duplicates == 3L,
  nrow(matches) == 2201L,
  sum(matches$Minutes) == 181081,
  sum(matches$Gls == 0) == 1063L,
  nrow(goals) == 1738L,
  sum(is.na(goals$Difficulty_Score)) == 3L,
  sum(is.na(goals$xg)) == 1255L,
  length(master_header) == 29L,
  length(derived) == 6L,
  ncol(matches) == 42L,
  uniqueN(matches$Comp) == 31L,
  uniqueN(matches$Venue) == 3L,
  matches[, .N, by = h2h_keys][N > 1, .N] == 0L,
  nrow(unmatched_goal_keys) == 3L,
  abs(b$famd_info$eig[1, 2] - 25.29678) < 1e-4,
  facts$stability_all == "TRUE",
  facts$stability_nopen == "TRUE",
  length(wave4$overall$differences) == 10000L,
  all(is.finite(wave4$overall$differences)),
  wave4$overall$interval[1] <= wave4$overall$interval[2],
  wave4$overall$probability_above_zero >= 0,
  wave4$overall$probability_above_zero <= 1,
  nrow(wave5_summary) == 14L,
  identical(unname(tools::md5sum(raw_csv)), unname(tools::md5sum(app_csv)))
)

tab <- data.frame(
  key = names(facts),
  value = vapply(facts, identity, character(1)),
  stringsAsFactors = FALSE
)
write.table(tab, out, sep = "\t", quote = FALSE, row.names = FALSE,
            col.names = TRUE, fileEncoding = "UTF-8")
'''


MASTER_GROUPS = [
    ("Identity and match", [
        ("Player", "text", "Player name.", "Separates the two careers."),
        ("Rk", "text", "FBref row number.", "Supports source auditing."),
        ("Date", "date", "Match date.", "Joins goals to match context and Elo."),
        ("Comp", "text", "Competition name.", "Identifies football context."),
        ("Round", "text", "Round or stage label.", "Builds Competition_Stage."),
        ("Venue", "text", "Home, Away, or Neutral.", "One FAMD input."),
        ("Squad", "text", "Source team label.", "Preserves the collected value."),
        ("Opponent", "text", "Source opponent label.", "Preserves the collected value."),
        ("Start", "text", "Whether the player started.", "Describes the appearance."),
    ]),
    ("Goal event", [
        ("Minute", "text", "Goal minute, including stoppage time.", "Orders multiple goals in a match."),
        ("Score", "text", "Score immediately after the goal.", "Describes match state."),
        ("Goalkeeper", "text", "Goalkeeper beaten.", "Event-level context."),
        ("Assist", "text", "Assisting player when recorded.", "Event-level context."),
        ("Notes", "text", "Goal note, such as Penalty kick.", "Creates the penalty flag."),
    ]),
    ("Clean joins and match context", [
        ("Opp_clean", "text", "Opponent without country prefix.", "Makes opponent joins reliable."),
        ("Squad_clean", "text", "Team without country prefix.", "Supports club and national labels."),
        ("Venue_match", "text", "Venue copied from the match log.", "Cross-checks the goal-log venue."),
        ("Result", "text", "Full-time match result.", "Adds match outcome context."),
        ("Minutes", "number", "Minutes in the appearance.", "Supplies the per-90 denominator."),
        ("Gls", "number", "Goals in that match, including zero.", "Retains scoreless appearances in match data."),
    ]),
    ("Difficulty, xG, and identifiers", [
        ("Opponent_Elo", "number", "Opponent rating or fallback value.", "Quantitative FAMD input."),
        ("Elo_Source", "text", "club, national, league_avg, or global_1500.", "Discloses how the rating was obtained."),
        ("Competition_Stage", "text", "Group, Knockout, Final, Qualifying, or Other.", "Categorical FAMD input."),
        ("Is_Away", "logical", "TRUE for an away match.", "Binary FAMD input."),
        ("Difficulty_Score", "number", "Oriented, standardized global FAMD Dim 1.", "Base difficulty index for weighting."),
        ("xg", "number", "Understat expected-goal value, when matched.", "Attached context; not used in the weight."),
        ("xg_source", "text", "understat or missing.", "Makes xG coverage explicit."),
        ("goal_id", "text", "Unique goal identifier.", "Tracks one row without ambiguity."),
        ("Weighted_Goal", "number", "Base value equal to Difficulty_Score.", "Stores the K=1 weighted index contribution."),
    ]),
]

DERIVED_GROUPS = [
    ("Dashboard-derived goal fields", [
        ("Season", "integer", "Season-ending year derived from Date.", "Supports time filters."),
        ("Era", "ordered factor", "Rise, Peak, Prime, Transition, or Late.", "Creates consistent career periods."),
        ("Penalty", "logical", "TRUE when Notes says Penalty kick.", "Runs included/excluded views."),
        ("Comp_Group", "text", "League, Continental, Domestic cup, International, or Other.", "Simplifies competition filters."),
        ("is_intl", "logical", "TRUE for national-team appearances.", "Separates club and international context."),
        ("Club_Goal", "logical", "TRUE when is_intl is FALSE.", "Builds the Overview club snapshot."),
    ])
]

MATCH_GROUPS = [
    ("Identity and schedule", [
        ("Player", "text", "Player name.", "Separates careers."),
        ("Date", "date", "Match date.", "Aligns matches, goals, and Elo."),
        ("Day", "text", "Day of week.", "Source context and auditing."),
        ("Comp", "text", "Competition.", "Groups football contexts."),
        ("Round", "text", "Round or matchweek.", "Builds stage categories."),
        ("Venue", "text", "Home, Away, or Neutral.", "FAMD input."),
        ("Result", "text", "Full-time result.", "Describes the appearance."),
        ("Squad", "text", "Source team label.", "Identifies the player's side."),
        ("Opponent", "text", "Source opponent label.", "Identifies the opposing side."),
        ("Start", "text", "Starter indicator.", "Describes playing role."),
        ("Pos", "text", "Recorded position.", "Appearance context; not in FAMD."),
        ("Min", "text", "Original minute value.", "Preserves the raw source field."),
    ]),
    ("Attacking match statistics", [
        ("Gls", "number", "Goals in the appearance; zero retained.", "Counts scoring and scoreless matches."),
        ("Ast", "number", "Assists.", "Context only; not in the current weight."),
        ("PK", "number", "Penalties scored.", "Source check for penalty context."),
        ("PKatt", "number", "Penalties attempted.", "Source check for penalty context."),
        ("Sh", "number", "Shots.", "Context only."),
        ("SoT", "number", "Shots on target.", "Context only."),
    ]),
    ("Discipline and actions", [
        ("CrdY", "number", "Yellow cards.", "Source context only."),
        ("CrdR", "number", "Red cards.", "Source context only."),
        ("Fls", "number", "Fouls committed.", "Source context only."),
        ("Fld", "number", "Fouls drawn.", "Source context only."),
        ("Off", "number", "Offsides.", "Source context only."),
        ("Crs", "number", "Crosses.", "Source context only."),
        ("TklW", "number", "Tackles won.", "Source context only."),
        ("Int", "number", "Interceptions.", "Source context only."),
        ("OG", "number", "Own goals.", "Source context only."),
        ("PKwon", "number", "Penalties won.", "Source context only."),
        ("PKcon", "number", "Penalties conceded.", "Source context only."),
        ("Match Report", "text", "FBref match-report link/value.", "Traceability to the source."),
    ]),
    ("Cleaned and analytical fields", [
        ("Season_End", "integer", "Season label from collection.", "Preserves source season grouping."),
        ("Squad_clean", "text", "Clean team name.", "Reliable joins and display."),
        ("Opp_clean", "text", "Clean opponent name.", "Reliable joins to Elo."),
        ("Minutes", "number", "Numeric minutes played.", "Complete per-90 denominator."),
        ("is_intl", "logical", "National-team match indicator.", "Chooses national versus club Elo."),
        ("Competition_Stage", "text", "Derived stage category.", "FAMD input."),
        ("Is_Away", "logical", "Away-match indicator.", "FAMD input."),
        ("Opponent_Elo", "number", "Opponent rating after fallback.", "FAMD input."),
        ("Elo_Source", "text", "Rating source or fallback.", "Audits rating coverage."),
        ("Difficulty_Score", "number", "Standardized global FAMD Dim 1.", "Joined to goals for weighting."),
        ("Season", "integer", "Season-ending year from Date.", "Dashboard time grouping."),
        ("Era", "ordered factor", "Five-part career era.", "Dashboard era grouping."),
    ]),
]


def extract_facts() -> dict[str, str]:
    """Read RDS-backed facts through a temporary plain Rscript."""
    with tempfile.TemporaryDirectory(prefix="statistics-guide-") as tmp:
        tmp_path = Path(tmp)
        r_path = tmp_path / "extract_statistics_guide_facts.R"
        facts_path = tmp_path / "facts.tsv"
        r_path.write_text(R_FACT_EXPORTER, encoding="utf-8")
        env = os.environ.copy()
        env["GUIDE_FACTS_PATH"] = str(facts_path)
        subprocess.run(
            ["Rscript", str(r_path)],
            cwd=ROOT,
            env=env,
            check=True,
            text=True,
            capture_output=True,
        )
        with facts_path.open(encoding="utf-8", newline="") as handle:
            return {row["key"]: row["value"] for row in csv.DictReader(handle, delimiter="\t")}


def n(facts: dict[str, str], key: str) -> int:
    return int(float(facts[key]))


def x(facts: dict[str, str], key: str) -> float:
    return float(facts[key])


def assert_close(actual: float, expected: float, tol: float = 5e-7) -> None:
    assert math.isclose(actual, expected, rel_tol=0, abs_tol=tol), (actual, expected)


def flatten(groups):
    return [row for _, rows in groups for row in rows]


def verify_facts(f: dict[str, str]) -> None:
    expected_ints = {
        "raw_goal_rows": 1738, "raw_match_rows": 2263, "raw_season_rows": 140,
        "raw_understat_rows": 487, "raw_transfer_rows": 20, "raw_bio_rows": 0,
        "club_elo_rows": 633265, "national_elo_rows": 1820,
        "header_junk": 12, "dnp_rows": 47, "duplicate_rows": 3,
        "valid_match_rows": 2201, "total_minutes": 181081,
        "scoreless_total": 1063, "goal_rows": 1738, "xg_matched": 483,
        "xg_missing": 1255, "difficulty_present": 1735, "difficulty_missing": 3,
        "master_field_count": 29, "derived_field_count": 6, "match_field_count": 42,
        "messi_goals": 847, "ronaldo_goals": 891,
        "messi_apps": 1031, "ronaldo_apps": 1170,
        "messi_minutes": 85131, "ronaldo_minutes": 95950,
        "messi_scoreless": 477, "ronaldo_scoreless": 586,
        "messi_penalties": 106, "ronaldo_penalties": 169,
        "messi_open_play": 741, "ronaldo_open_play": 722,
        "competition_count": 31, "venue_count": 3,
        "h2h_match_key_duplicates": 0, "h2h_unmatched_goal_keys": 3,
        "messi_h2h_scatter": 845, "ronaldo_h2h_scatter": 890,
        "champions_away_messi_apps": 81, "champions_away_ronaldo_apps": 92,
        "champions_away_messi_goals": 49, "champions_away_ronaldo_goals": 63,
        "world_cup_neutral_messi_apps": 29, "world_cup_neutral_ronaldo_apps": 25,
        "world_cup_neutral_messi_goals": 21, "world_cup_neutral_ronaldo_goals": 11,
        "mls_all_messi_apps": 76, "mls_all_ronaldo_apps": 0,
        "mls_all_messi_goals": 69, "mls_all_ronaldo_goals": 0,
        "fa_cup_all_messi_apps": 0, "fa_cup_all_ronaldo_apps": 1,
        "fa_cup_all_messi_goals": 0, "fa_cup_all_ronaldo_goals": 0,
        "mls_neutral_messi_apps": 0, "mls_neutral_ronaldo_apps": 0,
        "mls_neutral_messi_goals": 0, "mls_neutral_ronaldo_goals": 0,
        "wave5_summary_rows": 14, "raw_csv_rows": 1738,
        "raw_csv_columns": 29, "raw_csv_bytes": 431529,
    }
    for key, expected in expected_ints.items():
        assert n(f, key) == expected, (key, f[key], expected)
    assert f["stability_all"] == "TRUE" and f["stability_nopen"] == "TRUE"
    assert_close(x(f, "messi_raw90"), 0.8954435)
    assert_close(x(f, "ronaldo_raw90"), 0.8357478)
    assert_close(x(f, "messi_weighted90"), 0.1336718582)
    assert_close(x(f, "ronaldo_weighted90"), 0.0384164403)
    assert_close(x(f, "messi_h2h_cumulative"), 126.4402, 5e-5)
    assert_close(x(f, "ronaldo_h2h_cumulative"), 40.9562, 5e-5)
    assert_close(x(f, "messi_h2h_rolling30"), -0.0474, 5e-5)
    assert_close(x(f, "ronaldo_h2h_rolling30"), 0.0239, 5e-5)
    assert_close(x(f, "messi_penalty_share"), 106 / 847)
    assert_close(x(f, "ronaldo_penalty_share"), 169 / 891)
    assert_close(x(f, "famd_dim1_variance"), 25.29678, 1e-4)
    assert n(f, "wave4_reps") == 10000
    assert n(f, "wave4_finite") == 10000
    assert n(f, "wave4_seed") == 20260812
    assert_close(x(f, "wave4_gap"), 0.0952554179, 5e-10)
    assert_close(x(f, "wave4_ci_low"), -0.0119600327, 5e-10)
    assert_close(x(f, "wave4_ci_high"), 0.1999414634, 5e-10)
    assert_close(x(f, "wave4_probability"), 0.96, 5e-10)
    assert_close(x(f, "wave4_cohens_d"), 0.0899711631, 5e-10)
    assert_close(x(f, "summary_overall_messi_weighted90"), 0.1336718582)
    assert_close(x(f, "summary_overall_ronaldo_weighted90"), 0.0384164403)
    assert_close(x(f, "summary_overall_weighted_gap"), 0.0952554179)
    assert_close(x(f, "summary_overall_messi_raw90"), 0.8954434930)
    assert_close(x(f, "summary_overall_ronaldo_raw90"), 0.8357477853)
    assert f["raw_csv_md5"] == f["app_csv_md5"] == "c43c3f995b1f301b4328c846eab2cf27"
    assert f["master_fields"].split("||") == [r[0] for r in flatten(MASTER_GROUPS)]
    assert f["derived_fields"].split("||") == [r[0] for r in flatten(DERIVED_GROUPS)]
    assert f["match_fields"].split("||") == [r[0] for r in flatten(MATCH_GROUPS)]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=70, bottom=55, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_layout(table, widths, font_size=7.4) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    repeat_header(table.rows[0])
    for ridx, row in enumerate(table.rows):
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(font_size)
        if ridx == 0:
            for cell in row.cells:
                shade(cell, NAVY)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
        elif ridx % 2 == 0:
            for cell in row.cells:
                shade(cell, PALE_GRAY)


def add_table(doc, headers, rows, widths, font_size=7.8, style="Table Grid"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
    set_repeat_table_layout(table, widths, font_size)
    return table


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size, color in (("Title", 29, NAVY), ("Heading 1", 19, NAVY),
                              ("Heading 2", 12.5, MESSI), ("Heading 3", 10, RONALDO)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "MESSI vs RONALDO  •  STATISTICS EXPLAINED  •  WAVE 5"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(section.footer.paragraphs[0])


def add_page_title(doc, number, eyebrow, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{number:02d} / {TOTAL_PAGES}   {eyebrow.upper()}")
    r.font.name = "Aptos"
    r.font.size = Pt(7.8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.paragraph_format.space_after = Pt(7)
        for r in p.runs:
            r.font.size = Pt(9.2)
            r.font.color.rgb = RGBColor.from_string(MUTED)


def new_page(doc, number, eyebrow, title, subtitle=None):
    doc.add_page_break()
    add_page_title(doc, number, eyebrow, title, subtitle)


def add_callout(doc, label, text, fill=PALE_BLUE, accent=MESSI):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label.upper() + "  ")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(accent)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_formula(doc, formula, explanation):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, PALE_GRAY)
    set_cell_margins(cell, top=100, start=120, bottom=95, end=120)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(formula)
    r.font.name = "Cambria Math"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = cell.add_paragraph(explanation)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_field_table(doc, groups, include_group=True, max_rows=None):
    rows = []
    count = 0
    for group, fields in groups:
        for field in fields:
            if max_rows is not None and count >= max_rows:
                break
            rows.append((group if include_group else "",) + field)
            count += 1
    return add_table(
        doc,
        ["Group", "Field", "Type", "What it means", "Why we need it"],
        rows,
        [1.05, 1.14, 0.68, 2.15, 2.24],
        font_size=6.65,
    )


def fmt(v, digits=4):
    return f"{v:.{digits}f}"


def normalize_output_dashes(doc: Document) -> None:
    """Use ASCII hyphens in the exported study document."""
    replacements = str.maketrans({"\u2011": "-", "\u2012": "-", "\u2013": "-", "\u2014": "-", "\u2212": "-"})

    def normalize_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = run.text.translate(replacements)

    normalize_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                normalize_paragraphs(cell.paragraphs)
    for section in doc.sections:
        normalize_paragraphs(section.header.paragraphs)


def build_document(f: dict[str, str]) -> Document:
    doc = Document()
    configure_document(doc)

    # 1 — cover
    accent = doc.add_table(rows=1, cols=2)
    accent.autofit = False
    accent.cell(0, 0).width = Inches(3.6)
    accent.cell(0, 1).width = Inches(3.6)
    shade(accent.cell(0, 0), MESSI)
    shade(accent.cell(0, 1), RONALDO)
    for cell in accent.rows[0].cells:
        set_cell_margins(cell, top=35, bottom=35)
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MESSI vs RONALDO")
    r.font.name = "Aptos Display"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Statistics Explained")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A plain-English teaching guide through verified match-level uncertainty")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_callout(
        doc,
        "The question",
        "What changes when goals are compared per 90 minutes and then adjusted by a statistically derived match-context index?",
        PALE_GOLD,
        GOLD,
    )
    doc.add_heading("What this guide covers", level=2)
    p = doc.add_paragraph(
        "The collected sources, the cleaning path, every completed formula and metric, the current results, and the limits of the evidence through Dashboard Wave 5."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NEUTRAL • DATA-SNAPSHOT RESULTS • NOT A WINNER DECLARATION")
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(RONALDO)
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    add_table(doc, ["Snapshot", "Value"], [
        ("Goals", f"{n(f, 'goal_rows'):,}"),
        ("Valid appearances", f"{n(f, 'valid_match_rows'):,}"),
        ("Valid minutes", f"{n(f, 'total_minutes'):,}"),
        ("Coverage status", "Wave 5 complete and verified"),
    ], [2.3, 4.8], font_size=8.5)

    # 2 — contents and 60 seconds
    new_page(doc, 2, "Orientation", "Contents and the 60-second explanation")
    contents = [
        ("Project and data", "Sources; cleaning; architecture", "3–4"),
        ("Current snapshot", "Totals, minutes, penalties, raw and weighted rates", "5–6"),
        ("How context is built", "Elo fallback and global FAMD", "7–8"),
        ("How weighting works", "Equations, K sensitivity, density, deciles", "9–12"),
        ("Limits", "Coverage and interpretation caveats", "13"),
        ("Head-to-Head", "Trajectories, Elo scatter, penalties, filters", "14-16"),
        ("Inference", "Match bootstrap, interval, probability, d, subgroups", "17-19"),
        ("Field appendices", "29 master + 6 derived + 42 valid-match fields", "20-24"),
        ("Explain it aloud", "Glossary, speaking notes, stopping point", "25-26"),
        ("Wave 5 dashboard", "Methodology, definitive Summary, exact Raw Data", "27-29"),
    ]
    add_table(doc, ["Part", "What is inside", "Pages"], contents, [1.5, 4.9, 0.7], 8.2)
    doc.add_heading("The 60-second explanation", level=2)
    p = doc.add_paragraph(
        "We collected 1,738 goals and 2,201 valid appearances for Messi and Ronaldo. Every valid minute, including 1,063 scoreless appearances, stays in the per-90 denominator. One global FAMD combines opponent Elo, venue, competition stage, and an away indicator into a centred context index. The K control applies signed power. The Inference tab freezes the clicked K, penalty, competition, and venue settings, then independently resamples whole appearances within each player 10,000 times. It reports the full Messi-minus-Ronaldo gap distribution, a percentile interval, directional probability, and a conventional match-level Cohen's d without p-values or a winner flag."
    )
    p.paragraph_format.space_after = Pt(6)
    add_callout(doc, "Say it simply", "We compare scoring rates using all playing time, then stress-test a data-derived context index without pretending the index is a literal goal count.", PALE_TEAL, TEAL)
    doc.add_heading("How to read the numbers", level=2)
    add_bullets(doc, [
        "Raw goals per 90 are literal scoring rates.",
        "Weighted values are centred index units per 90, so they can be negative.",
        "The comparison is a fixed dataset snapshot, not live official career totals.",
    ])

    # 3 — sources
    new_page(doc, 3, "Data foundation", "What was collected—and what currently affects the statistics")
    source_rows = [
        ("FBref goal logs", f"{n(f,'raw_goal_rows'):,} goal rows", "One row per goal event", "Yes—goal numerators and goal context"),
        ("FBref match logs", f"{n(f,'raw_match_rows'):,} raw rows", "Appearances, minutes, scoreless matches", "Yes—FAMD and all per-90 denominators"),
        ("FBref season stats", f"{n(f,'raw_season_rows'):,} rows", "Season summaries and source checks", "No direct role in current weighting"),
        ("Understat", f"{n(f,'raw_understat_rows'):,} xG goal records; {n(f,'xg_matched'):,} matched", "Shot expected-goal context", "Attached, but not used in difficulty weight"),
        ("ClubElo", f"{n(f,'club_elo_rows'):,} lookup rows", "Historical club opponent ratings", "Yes—first Elo source for club matches"),
        ("National Elo", f"{n(f,'national_elo_rows'):,} lookup rows", "Historical national-team ratings", "Yes—for international matches"),
        ("Transfermarkt transfers", f"{n(f,'raw_transfer_rows'):,} rows", "Club-movement history", "No effect on current weighting"),
        ("Transfermarkt biography", f"{n(f,'raw_bio_rows'):,} usable rows", "Biographical details", "No effect on current weighting"),
    ]
    add_table(doc, ["Source", "Collected", "Why collected", "Current statistical role"], source_rows, [1.25, 1.55, 2.0, 2.3], 7.1)
    doc.add_heading("Coverage that must not be confused", level=2)
    add_bullets(doc, [
        f"Understat supplied xG for {n(f,'xg_matched'):,} of {n(f,'goal_rows'):,} goals; {n(f,'xg_missing'):,} goals lack xG.",
        "xG is descriptive enrichment only. The dashboard difficulty score does not use xG.",
        "Transfer rows can explain career journeys but do not enter FAMD, raw rates, weighted rates, or the K test.",
    ])
    add_callout(doc, "Say it simply", "Collected does not mean used. The current weight uses match context—not xG, transfer values, the player name, goal totals, or the player's own team Elo.")

    # 4 — cleaning
    new_page(doc, 4, "Transformation", "From raw match logs to the analysis population")
    cleaning = [
        ("Raw FBref match-log rows", f"{n(f,'raw_match_rows'):,}", "Starting snapshot"),
        ("Repeated header rows removed", f"−{n(f,'header_junk'):,}", "Rows such as Gls = 'Gls'"),
        ("No-minute / DNP rows removed", f"−{n(f,'dnp_rows'):,}", "Not valid played appearances"),
        ("Exact Nations League duplicates removed", f"−{n(f,'duplicate_rows'):,}", "Same player-date-opponent-context key"),
        ("Valid appearances retained", f"{n(f,'valid_match_rows'):,}", "The FAMD and per-90 population"),
    ]
    add_table(doc, ["Cleaning step", "Rows", "Reason"], cleaning, [3.2, 1.0, 2.9], 8.2)
    doc.add_heading("The non-negotiable denominator", level=2)
    p = doc.add_paragraph(
        f"The valid set contains {n(f,'scoreless_total'):,} scoreless appearances ({100*n(f,'scoreless_total')/n(f,'valid_match_rows'):.1f}%). They stay in the analysis. Removing them would compare only matches in which a goal happened and would exaggerate both per-90 rates."
    )
    flow_rows = [
        ("2,201 valid appearances", "One global FAMD on all valid matches", "A match difficulty score"),
        ("1,738 goal events", "Join the match score to each goal", "1,735 scored goals + 3 missing"),
        ("All 181,081 valid minutes", "Stay in player denominators", "Raw/90 and weighted index/90"),
    ]
    add_table(doc, ["Input", "Transformation", "Output"], flow_rows, [2.1, 2.8, 2.2], 8.1)
    add_callout(doc, "Say it simply", "A player does not vanish from the denominator merely because he did not score.", PALE_GOLD, GOLD)

    # 5 — baseline
    new_page(doc, 5, "Current results", "The verified baseline snapshot", "Baseline means K = 1 with penalties included unless stated otherwise.")
    baseline = [
        ("Goals", f"{n(f,'messi_goals'):,}", f"{n(f,'ronaldo_goals'):,}"),
        ("Valid appearances", f"{n(f,'messi_apps'):,}", f"{n(f,'ronaldo_apps'):,}"),
        ("Valid minutes", f"{n(f,'messi_minutes'):,}", f"{n(f,'ronaldo_minutes'):,}"),
        ("Scoreless appearances", f"{n(f,'messi_scoreless'):,}", f"{n(f,'ronaldo_scoreless'):,}"),
        ("Penalty goals", f"{n(f,'messi_penalties'):,}", f"{n(f,'ronaldo_penalties'):,}"),
        ("Raw goals / 90", fmt(x(f,'messi_raw90')), fmt(x(f,'ronaldo_raw90'))),
        ("Weighted index / 90", fmt(x(f,'messi_weighted90')), fmt(x(f,'ronaldo_weighted90'))),
    ]
    table = add_table(doc, ["Metric", "Messi", "Ronaldo"], baseline, [3.4, 1.85, 1.85], 8.6)
    shade(table.rows[0].cells[1], MESSI); shade(table.rows[0].cells[2], RONALDO)
    doc.add_heading("Penalty-included versus penalty-excluded", level=2)
    penalty = [
        ("Included", "847", fmt(x(f,'messi_raw90')), fmt(x(f,'messi_weighted90')), "891", fmt(x(f,'ronaldo_raw90')), fmt(x(f,'ronaldo_weighted90'))),
        ("Excluded", "741", fmt(x(f,'messi_raw90_nopen')), fmt(x(f,'messi_weighted90_nopen')), "722", fmt(x(f,'ronaldo_raw90_nopen')), fmt(x(f,'ronaldo_weighted90_nopen'))),
    ]
    add_table(doc, ["Penalties", "M goals", "M raw/90", "M weighted/90", "R goals", "R raw/90", "R weighted/90"], penalty, [0.85,0.65,0.9,1.05,0.65,0.9,1.05], 6.9)
    doc.add_heading("Overview club snapshot counts", level=2)
    club_rows = [
        ("Messi", "Barcelona", n(f,"club_messi_barcelona")), ("Messi", "Paris Saint-Germain", n(f,"club_messi_psg")),
        ("Messi", "Inter Miami", n(f,"club_messi_inter_miami")), ("Ronaldo", "Sporting CP", n(f,"club_ronaldo_sporting")),
        ("Ronaldo", "Manchester United", n(f,"club_ronaldo_man_utd")), ("Ronaldo", "Real Madrid", n(f,"club_ronaldo_real_madrid")),
        ("Ronaldo", "Juventus", n(f,"club_ronaldo_juventus")), ("Ronaldo", "Al-Nassr", n(f,"club_ronaldo_al_nassr")),
    ]
    add_table(doc, ["Player", "Club", "Goals in this dataset"], club_rows, [1.25,3.65,2.2], 7.5)

    # 6 — raw metrics
    new_page(doc, 6, "Calculation", "Totals, rates, scoreless appearances, and penalties")
    add_formula(doc, "Raw/90 = 90 × goals ÷ total valid minutes", "A literal scoring rate: how many goals the observed pace would produce in 90 minutes.")
    add_table(doc, ["Symbol", "Meaning", "Rule"], [
        ("goals", "Eligible goal rows for the player", "Includes penalties unless the toggle excludes them"),
        ("total valid minutes", "Minutes across every valid appearance", "Never restricted to goal-scoring matches"),
        ("90", "Standard match-length scale", "Makes players with different minutes comparable"),
    ], [1.2,3.0,2.9], 8.0)
    doc.add_heading("Worked examples", level=2)
    add_bullets(doc, [
        f"Messi: 90 × {n(f,'messi_goals'):,} ÷ {n(f,'messi_minutes'):,} = {fmt(x(f,'messi_raw90'))} raw goals per 90.",
        f"Ronaldo: 90 × {n(f,'ronaldo_goals'):,} ÷ {n(f,'ronaldo_minutes'):,} = {fmt(x(f,'ronaldo_raw90'))} raw goals per 90.",
        f"Without penalties, goal numerators become 741 and 722, but the minute denominators remain {n(f,'messi_minutes'):,} and {n(f,'ronaldo_minutes'):,}.",
    ])
    doc.add_heading("What each completed count means", level=2)
    count_rows = [
        ("Total goals", "Number of goal rows", "Literal count; no context adjustment"),
        ("Appearances", "Valid played match rows", "Includes scoreless appearances"),
        ("Minutes", "Sum of numeric played minutes", "Denominator for every per-90 view"),
        ("Scoreless", "Valid appearances where Gls = 0", "Prevents scorer-only selection bias"),
        ("Penalties", "Goal rows marked Penalty kick", "Optional numerator exclusion"),
    ]
    add_table(doc, ["Metric", "How calculated", "Interpretation"], count_rows, [1.2,2.7,3.2], 8.0)
    add_callout(doc, "Say it simply", "The penalty switch removes penalty goals, not the matches or minutes in which they occurred.")

    # 7 — Elo
    new_page(doc, 7, "Calculation", "How every valid match receives an opponent Elo")
    elo_steps = [
        ("1", "Club Elo by opponent and date", n(f,"match_elo_club"), "Used for matched club appearances"),
        ("2", "National Elo by team and date", n(f,"match_elo_national"), "Used for matched internationals"),
        ("3", "Competition-average Elo", n(f,"match_elo_league_avg"), "Fallback when a direct rating is absent"),
        ("4", "Global 1500", n(f,"match_elo_global_1500"), "Last fallback; no match is dropped"),
    ]
    add_table(doc, ["Order", "Source / fallback", "Appearances", "Purpose"], elo_steps, [0.55,2.45,1.05,3.05], 8.0)
    direct = n(f,"match_elo_direct")
    dol = n(f,"match_elo_direct_or_league")
    doc.add_heading("Coverage", level=2)
    add_bullets(doc, [
        f"Direct club or national Elo: {direct:,} of {n(f,'valid_match_rows'):,} appearances ({100*direct/n(f,'valid_match_rows'):.1f}%).",
        f"Direct or competition-average: {dol:,} of {n(f,'valid_match_rows'):,} ({100*dol/n(f,'valid_match_rows'):.1f}%).",
        f"The remaining {n(f,'match_elo_global_1500'):,} appearances use 1500. All valid appearances therefore enter FAMD.",
    ])
    doc.add_heading("What Elo does—and does not do", level=2)
    add_bullets(doc, [
        "It supplies a time-specific estimate of opponent strength.",
        "It is only one of four FAMD inputs; it is not itself the final weight.",
        "The player's own team Elo is deliberately excluded, avoiding a penalty for playing in a stronger side or era.",
        "Fallback values preserve coverage but are less specific than direct historical ratings.",
    ])
    add_callout(doc, "Say it simply", "First use the opponent's historical rating; if it is missing, use the best documented fallback and keep the match.", PALE_GOLD, GOLD)

    # 8 — FAMD
    new_page(doc, 8, "Method", "One global FAMD turns mixed match context into a common index")
    add_formula(doc, "Dᵢ = standardized FAMD Dim 1 score", "The same global model is fitted once to all 2,201 valid appearances for both players.")
    add_table(doc, ["Input", "Kind", "Dim 1 contribution", "Role"], [
        ("Opponent_Elo", "continuous", f"{x(f,'famd_contrib_opponent_elo'):.1f}%", "Opponent strength"),
        ("Venue", "categorical", f"{x(f,'famd_contrib_venue'):.1f}%", "Home / Away / Neutral"),
        ("Competition_Stage", "categorical", f"{x(f,'famd_contrib_competition_stage'):.1f}%", "Group / knockout / final / qualifying / other"),
        ("Is_Away", "categorical", f"{x(f,'famd_contrib_is_away'):.1f}%", "Explicit away indicator"),
    ], [1.7,1.1,1.35,2.95], 7.7)
    doc.add_heading("Orientation and standardization", level=2)
    add_bullets(doc, [
        "FAMD is used because the inputs mix one continuous variable with categorical variables; PCA alone is not appropriate for that mixture.",
        "Dim 1 is flipped when necessary so its correlation with opponent Elo is positive: higher is intended to mean harder.",
        "Dim 1 is then centred and scaled across valid appearances. Zero is the centre of this index, not 'no difficulty'.",
        f"The first dimension explains {x(f,'famd_dim1_variance'):.1f}% of total variation. Most variation is therefore outside this one-number summary.",
    ])
    add_callout(doc, "Major caution", "Venue (48.6%) and Is_Away (48.3%) overlap conceptually. Together they dominate Dim 1, while Opponent Elo contributes only 0.8%. The current index is largely an away/home axis—not a pure opponent-strength measure.", PALE_RED, RONALDO)

    # 9 — weighting formulas
    new_page(doc, 9, "Calculation", "From a match index to a selected-K weighted rate")
    add_formula(doc, "WGᵢ = Dᵢ", "Base weighted goal at K = 1. Each scored goal inherits its match's difficulty index.")
    add_formula(doc, "WGᵢ(K) = sign(Dᵢ) × |Dᵢ|ᴷ", "Signed power preserves the sign and remains finite for negative scores at fractional K.")
    add_formula(doc, "Weighted/90 = 90 × Σ WGᵢ(K) ÷ Σ minutes", "Goal-index contributions form the numerator; all valid appearance minutes form the denominator.")
    add_formula(doc, "Gap(K) = Messi Weighted/90 − Ronaldo Weighted/90", "A positive gap means Messi's index rate is higher at that tested K; a negative gap would mean Ronaldo's is higher.")
    add_table(doc, ["Symbol", "Meaning"], [
        ("i", "A goal with an available match difficulty score"),
        ("Dᵢ", "Oriented, standardized FAMD Dim 1 score"),
        ("K", "Sensitivity exponent from 0.5 to 3.0"),
        ("sign(Dᵢ)", "+1 for positive scores, −1 for negative scores"),
        ("Σ minutes", "All valid player minutes, including scoreless appearances"),
    ], [1.35,5.75], 8.0)
    doc.add_heading("Worked signed-power example", level=2)
    p = doc.add_paragraph("If Dᵢ = −0.64 and K = 0.5, then WGᵢ(K) = −1 × |−0.64|⁰·⁵ = −0.80. Plain (−0.64)⁰·⁵ would be invalid in ordinary real-number calculation; signed power is why the dashboard remains finite.")

    # 10 — K included
    new_page(doc, 10, "Weighting Lab", "K sensitivity with penalties included")
    rows = []
    for kval, key in [(0.5,"0_5"),(1.0,"1_0"),(1.5,"1_5"),(3.0,"3_0")]:
        rows.append((f"{kval:.1f}", fmt(x(f,f"k_{key}_messi_all")), fmt(x(f,f"k_{key}_ronaldo_all")), f"{x(f,f'k_{key}_gap_all'):+.4f}"))
    add_table(doc, ["K", "Messi weighted/90", "Ronaldo weighted/90", "Messi − Ronaldo gap"], rows, [0.8,2.05,2.05,2.2], 8.8)
    doc.add_heading("How to read the pattern", level=2)
    add_bullets(doc, [
        "K below 1 compresses large absolute scores and lifts smaller absolute values relative to them.",
        "K = 1 preserves the base difficulty index.",
        "K above 1 amplifies large absolute scores and shrinks values whose absolute size is below 1.",
        f"At K = 3, both rates can cross below zero: Messi {fmt(x(f,'k_3_0_messi_all'))} and Ronaldo {fmt(x(f,'k_3_0_ronaldo_all'))}. These are valid centred-index outputs.",
    ])
    add_callout(doc, "Verified stability result", f"No lead change occurs at any of the 26 tested values from K = 0.5 to 3.0. The positive gap ranges from {fmt(x(f,'gap_min_all'))} to {fmt(x(f,'gap_max_all'))} index units per 90.", PALE_TEAL, TEAL)
    doc.add_heading("What the test does not establish", level=2)
    p = doc.add_paragraph("It does not attach sampling uncertainty, prove causation, or show that the FAMD inputs are the only valid definition of difficulty. It asks a narrower question: does the ordering change when the exponent changes on this fixed dataset and index?")

    # 11 — K excluded
    new_page(doc, 11, "Weighting Lab", "K sensitivity with penalties excluded")
    rows = []
    for kval, key in [(0.5,"0_5"),(1.0,"1_0"),(1.5,"1_5"),(3.0,"3_0")]:
        rows.append((f"{kval:.1f}", fmt(x(f,f"k_{key}_messi_nopen")), fmt(x(f,f"k_{key}_ronaldo_nopen")), f"{x(f,f'k_{key}_gap_nopen'):+.4f}"))
    add_table(doc, ["K", "Messi weighted/90", "Ronaldo weighted/90", "Messi − Ronaldo gap"], rows, [0.8,2.05,2.05,2.2], 8.8)
    doc.add_heading("The denominator contract", level=2)
    add_bullets(doc, [
        f"Eligible goal numerators become 741 for Messi and 722 for Ronaldo.",
        f"Minute denominators remain {n(f,'messi_minutes'):,} and {n(f,'ronaldo_minutes'):,}; the toggle never removes appearances.",
        f"At K = 3 the valid rates are {fmt(x(f,'k_3_0_messi_nopen'))} and {fmt(x(f,'k_3_0_ronaldo_nopen'))}, not errors.",
    ])
    add_callout(doc, "Verified stability result", f"No lead change occurs at any of the same 26 K values without penalties. The positive gap ranges from {fmt(x(f,'gap_min_nopen'))} to {fmt(x(f,'gap_max_nopen'))} index units per 90.", PALE_TEAL, TEAL)
    doc.add_heading("Lead stability calculation", level=2)
    p = doc.add_paragraph("For each penalty setting, the dashboard computes Gap(K) at K = 0.5, 0.6, …, 3.0. It checks whether adjacent gaps touch zero or change sign. Here all 26 gaps are positive in both settings, so the label is “No lead change across K=0.5–3.0.”")
    add_callout(doc, "Say it simply", "The ordering survives this exponent stress test, but that is not the same as proving a population-level difference.")

    # 12 — density and deciles
    new_page(doc, 12, "Weighting Lab", "Base-score density, fixed deciles, counts, and shares")
    add_table(doc, ["Penalty setting", "Messi median", "Ronaldo median", "Messi D10", "Ronaldo D10"], [
        ("Included", f"{x(f,'messi_median_all'):.3f}", f"{x(f,'ronaldo_median_all'):.3f}", f"{n(f,'messi_d10_all')} ({100*x(f,'messi_d10_share_all'):.1f}%)", f"{n(f,'ronaldo_d10_all')} ({100*x(f,'ronaldo_d10_share_all'):.1f}%)"),
        ("Excluded", f"{x(f,'messi_median_nopen'):.3f}", f"{x(f,'ronaldo_median_nopen'):.3f}", f"{n(f,'messi_d10_nopen')} ({100*x(f,'messi_d10_share_nopen'):.1f}%)", f"{n(f,'ronaldo_d10_nopen')} ({100*x(f,'ronaldo_d10_share_nopen'):.1f}%)"),
    ], [1.35,1.35,1.35,1.5,1.5], 7.8)
    doc.add_heading("Normalized density", level=2)
    add_bullets(doc, [
        "A smooth curve estimates the distribution of K=1 difficulty scores for each player's eligible goals.",
        "Each player's curve is normalized to area 1. Curve height is relative concentration, not a goal count.",
        "The density stays on the base score; moving K does not change it. The penalty toggle changes which goals contribute.",
    ])
    doc.add_heading("Fixed global deciles", level=2)
    add_bullets(doc, [
        f"The {n(f,'difficulty_present'):,} goals with a score are ranked together before penalty filtering.",
        "Ranks are split into D1 (easiest) through D10 (hardest). Ties use a stable first-occurrence rule.",
        f"The D10 base-score range is {x(f,'d10_min'):.3f} to {x(f,'d10_max'):.3f} and does not move when penalties are excluded.",
        "Count = eligible goals in a decile. Player share = that count ÷ all eligible scored goals for the same player.",
    ])
    add_formula(doc, "Player share in D10 = D10 eligible goals ÷ all eligible scored goals", f"Included example: Messi = {n(f,'messi_d10_all')} ÷ {n(f,'messi_eligible_all')} = {100*x(f,'messi_d10_share_all'):.1f}%; Ronaldo = {n(f,'ronaldo_d10_all')} ÷ {n(f,'ronaldo_eligible_all')} = {100*x(f,'ronaldo_d10_share_all'):.1f}%.")
    add_callout(doc, "Say it simply", "Density compares distribution shape; decile bars compare counts. They answer different questions.")

    # 13 — caveats
    new_page(doc, 13, "Interpretation", "Caveats deserve the same prominence as the results")
    caveats = [
        ("Three missing difficulty scores", f"Raw rates include all {n(f,'goal_rows'):,} goals. Weighted views use {n(f,'difficulty_present'):,}; 2 Messi goals and 1 Ronaldo goal lack match context."),
        ("xG coverage", f"{n(f,'xg_missing'):,} goals lack xG. xG is attached but does not enter the current weight."),
        ("One dimension is partial", f"Dim 1 explains {x(f,'famd_dim1_variance'):.1f}% of total FAMD variation; 74.7% is not represented by this index."),
        ("Away/home dominance", "Venue contributes 48.6% and Is_Away 48.3%; their overlap makes Dim 1 largely an away/home axis."),
        ("Opponent Elo is small", "Opponent Elo contributes 0.8% to Dim 1, so the score should not be described as mainly opponent strength."),
        ("Negative values are valid", "They are centred index values, not goals being taken away from the historical record."),
        ("K stability is narrow", "It is a sensitivity check on one transformation, not a causal estimate or winner declaration."),
        ("Bootstrap scope is conditional", "The interval describes match-resampling uncertainty under the clicked K, penalty, competition, venue, dataset, and weighting definition."),
        ("Snapshot limits", "The data are a collected snapshot and may differ from live official career totals or alternative providers."),
    ]
    add_table(doc, ["Caveat", "What it means"], caveats, [2.0,5.1], 7.7)
    add_callout(doc, "Responsible conclusion", "Wave 4 quantifies match-resampling uncertainty under the frozen dashboard assumptions. It does not settle every football definition of difficulty, remove model limitations, or declare a final winner.", PALE_RED, RONALDO)
    doc.add_heading("Best neutral wording", level=2)
    p = doc.add_paragraph(f"“At the default settings, the observed Messi-minus-Ronaldo weighted-index gap is {x(f,'wave4_gap'):.4f} per 90. The seeded 95% match-bootstrap interval is [{x(f,'wave4_ci_low'):.4f}, {x(f,'wave4_ci_high'):.4f}], and {100*x(f,'wave4_probability'):.1f}% of resampled gaps are above zero. These values are conditional on the dataset and index—not a final winner declaration.”")

    # 14 - Head-to-Head trajectories
    new_page(doc, 14, "Head-to-Head", "How the career trajectories are aligned and calculated", "Default view: K = 1, penalties included, all competitions, all venues, age axis.")
    add_formula(doc, "Cumulative at t = Σ selected-K contributions through appearance t", "The curve adds each eligible goal's signed-power contribution in filtered appearance order. Scoreless appearances contribute zero and stay on the timeline.")
    add_formula(doc, "Rolling-30/90 = 90 × Σ contributions in latest 30 appearances ÷ Σ minutes in those 30", "The first 29 filtered appearances have no rolling value because a full 30-appearance window does not yet exist.")
    add_table(doc, ["Default endpoint", "Messi", "Ronaldo"], [
        ("Cumulative weighted index", fmt(x(f,"messi_h2h_cumulative")), fmt(x(f,"ronaldo_h2h_cumulative"))),
        ("Latest rolling-30 index / 90", fmt(x(f,"messi_h2h_rolling30")), fmt(x(f,"ronaldo_h2h_rolling30"))),
    ], [3.3, 1.9, 1.9], 8.6)
    doc.add_heading("Age versus calendar date", level=2)
    add_bullets(doc, [
        "Age aligns the careers by life stage rather than by the year on the calendar.",
        "Age is derived from fixed display metadata: Messi 1987-06-24; Ronaldo 1985-02-05.",
        "The age-30 guide appears only on the age axis. Calendar date preserves the actual historical sequence.",
        "Changing competition or venue restarts the cumulative line, appearance count, and rolling window inside that selected scope.",
    ])
    add_callout(doc, "Interpret carefully", "Cumulative values answer how contributions accumulated in the selected scope. Rolling values answer what the latest 30 filtered appearances looked like. They are different questions, and neither is a confidence interval.", PALE_GOLD, GOLD)

    # 15 - opponent Elo scatter
    new_page(doc, 15, "Head-to-Head", "Opponent Elo scatter and descriptive LOESS")
    add_formula(doc, "Marker y = sign(Difficulty Score) × |Difficulty Score|ᴷ", "Each marker is one eligible goal. The x-axis is the actual unbinned opponent Elo; no jitter is applied.")
    add_table(doc, ["Player", "Eligible goal markers", "Distinct Elo values", "Default cumulative contribution"], [
        ("Messi", f"{n(f,'messi_h2h_scatter'):,}", f"{n(f,'messi_h2h_distinct_elo'):,}", fmt(x(f,"messi_h2h_cumulative"))),
        ("Ronaldo", f"{n(f,'ronaldo_h2h_scatter'):,}", f"{n(f,'ronaldo_h2h_distinct_elo'):,}", fmt(x(f,"ronaldo_h2h_cumulative"))),
    ], [1.2, 1.8, 1.55, 2.55], 8.0)
    doc.add_heading("What the chart can show", level=2)
    add_bullets(doc, [
        "Horizontal position shows the opponent rating attached to that goal's match.",
        "Vertical position shows the selected-K index contribution, which may be negative because the score is centred.",
        "The tooltip retains player, date, opponent, competition, venue, Elo, base score, selected-K contribution, and penalty status.",
        "A player-specific LOESS curve summarizes the local shape using the actual marker values.",
    ])
    doc.add_heading("What LOESS does not show", level=2)
    add_bullets(doc, [
        "It is a descriptive smoother, not a causal model, significance test, or prediction interval.",
        "It has no confidence band in the Head-to-Head view.",
        "It is suppressed when fewer than 10 goals or 5 distinct Elo values remain, because sparse smoothers are misleading.",
    ])
    add_callout(doc, "Say it simply", "The dots are the observed eligible goals; the curve is only a visual summary of their pattern.", PALE_TEAL, TEAL)

    # 16 - penalty dependency and filters
    new_page(doc, 16, "Head-to-Head", "Penalty dependency, filters, and sparse-scope rules")
    add_table(doc, ["Player", "Penalty goals", "Open-play goals", "Penalty share"], [
        ("Messi", f"{n(f,'messi_penalties'):,}", f"{n(f,'messi_open_play'):,}", f"{100*x(f,'messi_penalty_share'):.1f}%"),
        ("Ronaldo", f"{n(f,'ronaldo_penalties'):,}", f"{n(f,'ronaldo_open_play'):,}", f"{100*x(f,'ronaldo_penalty_share'):.1f}%"),
    ], [1.5, 1.7, 1.8, 2.1], 8.5)
    add_bullets(doc, [
        "The composition chart always discloses both goal types. When penalties are excluded, the penalty segment is muted and labeled excluded rather than hidden.",
        "Penalty exclusion changes goal contributions only; it never removes valid appearances or minutes.",
        f"Weighted coverage is {n(f,'difficulty_present'):,} of {n(f,'goal_rows'):,} goals ({100*n(f,'difficulty_present')/n(f,'goal_rows'):.1f}%). Missing-score goals remain in raw composition counts.",
        f"The shared filters offer {n(f,'competition_count'):,} competition labels plus All competitions, and Home, Away, Neutral, plus All venues.",
    ])
    doc.add_heading("Verified filter examples", level=2)
    add_table(doc, ["Scope", "Messi", "Ronaldo", "Why it matters"], [
        ("Champions Lg + Away", f"{n(f,'champions_away_messi_apps')} apps / {n(f,'champions_away_messi_goals')} goals", f"{n(f,'champions_away_ronaldo_apps')} apps / {n(f,'champions_away_ronaldo_goals')} goals", "Both players remain"),
        ("World Cup + Neutral", f"{n(f,'world_cup_neutral_messi_apps')} apps / {n(f,'world_cup_neutral_messi_goals')} goals", f"{n(f,'world_cup_neutral_ronaldo_apps')} apps / {n(f,'world_cup_neutral_ronaldo_goals')} goals", "Neutral international scope"),
        ("MLS + All", f"{n(f,'mls_all_messi_apps')} apps / {n(f,'mls_all_messi_goals')} goals", f"{n(f,'mls_all_ronaldo_apps')} apps", "Ronaldo is N/A"),
        ("FA Cup + All", f"{n(f,'fa_cup_all_messi_apps')} apps", f"{n(f,'fa_cup_all_ronaldo_apps')} app / {n(f,'fa_cup_all_ronaldo_goals')} goals", "Messi N/A; Ronaldo 0.0000"),
        ("MLS + Neutral", f"{n(f,'mls_neutral_messi_apps')} apps", f"{n(f,'mls_neutral_ronaldo_apps')} apps", "Explanatory empty state"),
    ], [1.7, 1.45, 1.45, 2.5], 7.2)
    add_callout(doc, "N/A versus zero", "N/A means a player has no valid appearances in the selected scope. 0.0000 means appearances exist but no eligible weighted contribution remains. Keeping those states separate prevents a missing comparison from looking like a real zero.", PALE_GOLD, GOLD)
    add_callout(doc, "Wave 4 handoff", "Head-to-Head remains descriptive and updates immediately. Inference changes only when Update analysis is clicked, then displays the frozen settings and warns if the live controls change.", PALE_TEAL, TEAL)

    # 17 - bootstrap method and overall results
    new_page(doc, 17, "Inference", "How the match-level bootstrap measures uncertainty", "Default frozen scope: K = 1, penalties included, all competitions, all venues.")
    add_formula(doc, "Rate* = 90 × Σ resampled match contributions ÷ Σ resampled match minutes", "Each player is resampled independently with replacement, using exactly that player's selected-scope appearance count.")
    add_formula(doc, "Gap* = Messi Rate* - Ronaldo Rate*", "Repeat 10,000 times. The 2.5th and 97.5th percentiles form the displayed 95% interval.")
    add_formula(doc, "P(Gap* > 0) = count(Gap* > 0) ÷ 10,000", "A directional share of the bootstrap distribution, not a p-value or binary decision.")
    add_table(doc, ["Default result", "Verified value", "How to read it"], [
        ("Observed ratio-of-sums gap / 90", f"{x(f,'wave4_gap'):+.4f}", "Messi - Ronaldo under frozen settings"),
        ("95% percentile interval", f"[{x(f,'wave4_ci_low'):+.4f}, {x(f,'wave4_ci_high'):+.4f}]", "Middle 95% of seeded resampled gaps"),
        ("P(gap > 0)", f"{x(f,'wave4_probability'):.3f}", "Share of resampled gaps above zero"),
        ("Match-level Cohen's d", f"{x(f,'wave4_cohens_d'):+.3f}", "Pooled-SD appearance-rate difference"),
        ("Replicates / seed", f"{n(f,'wave4_reps'):,} / {n(f,'wave4_seed')}", "Deterministic and fully finite"),
    ], [2.15, 1.6, 3.35], 7.8)
    add_callout(doc, "Interval interpretation", "The default percentile interval crosses zero. Report that numerical fact and the full distribution; do not convert it into a winner, significant/not-significant badge, or p-value.", PALE_GOLD, GOLD)

    # 18 - era subgroups
    new_page(doc, 18, "Inference", "Career-era subgroup results", "Every row inherits the clicked global competition and venue scope.")
    era_specs = [
        ("Rise (2002-08)", "rise_2002_08"), ("Peak (2009-14)", "peak_2009_14"),
        ("Prime (2015-18)", "prime_2015_18"), ("Transition (2019-22)", "transition_2019_22"),
        ("Late (2023-26)", "late_2023_26"),
    ]
    era_rows = []
    for label, slug in era_specs:
        era_rows.append((
            label,
            f"{n(f,f'era_{slug}_messi_apps')} / {x(f,f'era_{slug}_messi_rate'):.4f}",
            f"{n(f,f'era_{slug}_ronaldo_apps')} / {x(f,f'era_{slug}_ronaldo_rate'):.4f}",
            f"{x(f,f'era_{slug}_gap'):+.4f}",
            f"[{x(f,f'era_{slug}_ci_low'):+.4f}, {x(f,f'era_{slug}_ci_high'):+.4f}]",
            f"{x(f,f'era_{slug}_cohens_d'):+.3f}",
        ))
    add_table(doc, ["Era", "Messi apps / rate", "Ronaldo apps / rate", "Gap", "95% interval", "d"], era_rows, [1.3,1.35,1.35,.75,1.65,.7], 6.7)
    add_bullets(doc, [
        "Rates are descriptive whenever appearances exist.",
        "Intervals and Cohen's d require at least two appearances for both players.",
        "A row is marked sparse in the dashboard when either player has fewer than 30 appearances.",
        "Era comparisons are exploratory; they are not adjusted for repeated subgroup inspection.",
    ])
    add_callout(doc, "Subgroup caution", "Differences across eras can reflect career timing, teams, competitions, and the index definition. The table quantifies conditional patterns; it does not isolate a causal era effect.", PALE_RED, RONALDO)

    # 19 - competition families and effect size
    new_page(doc, 19, "Inference", "Competition-family results and Cohen's d")
    family_specs = [
        ("League", "league"), ("Continental", "continental"),
        ("Domestic cup", "domestic_cup"), ("International", "international"),
        ("Other", "other"),
    ]
    family_rows = []
    for label, slug in family_specs:
        family_rows.append((
            label,
            f"{n(f,f'family_{slug}_messi_apps')} / {x(f,f'family_{slug}_messi_rate'):.4f}",
            f"{n(f,f'family_{slug}_ronaldo_apps')} / {x(f,f'family_{slug}_ronaldo_rate'):.4f}",
            f"{x(f,f'family_{slug}_gap'):+.4f}",
            f"[{x(f,f'family_{slug}_ci_low'):+.4f}, {x(f,f'family_{slug}_ci_high'):+.4f}]",
            f"{x(f,f'family_{slug}_cohens_d'):+.3f}",
            "Sparse" if f[f"family_{slug}_sparse"] == "TRUE" else "Adequate count",
        ))
    add_table(doc, ["Family", "Messi apps / rate", "Ronaldo apps / rate", "Gap", "95% interval", "d", "Count"], family_rows, [1.05,1.3,1.3,.7,1.5,.55,.75], 6.25)
    add_formula(doc, "d = (mean Messi appearance rate - mean Ronaldo appearance rate) ÷ pooled SD", "This conventional match-level effect size uses appearance-level weighted rates. It is not the primary ratio-of-sums gap and does not use the bootstrap means.")
    add_bullets(doc, [
        "Domestic cup and Other are sparse at the default scope because at least one player has fewer than 30 appearances.",
        "If an exact competition is selected, only its inherited competition family may remain.",
        "Scoreless appearances stay in both family denominators; unmapped scoreless-only competition labels fall into Other.",
        "N/A is distinct from zero: no appearances suppress a comparison, while real appearances with no eligible contribution can produce 0.0000.",
    ])

    # 20 - master fields A
    new_page(doc, 20, "Field appendix A", "Master goal table: fields 1–15 of 29", "File: data/processed/goals_master_final.csv • one row per goal")
    master_rows = flatten(MASTER_GROUPS)
    # Preserve group labels while slicing.
    rows_a = []
    for group, fields in MASTER_GROUPS:
        for field in fields:
            rows_a.append((group,) + field)
    add_table(doc, ["Group", "Field", "Type", "What it means", "Why we need it"], rows_a[:15], [1.05,1.14,0.68,2.15,2.24], 6.65)
    add_callout(doc, "Appendix reading rule", "A field can be present for traceability without entering the current formula. The final column says whether its purpose is analysis, joining, display, or auditing.")

    # 21 - master fields B
    new_page(doc, 21, "Field appendix A", "Master goal table: fields 16–29 of 29")
    add_table(doc, ["Group", "Field", "Type", "What it means", "Why we need it"], rows_a[15:], [1.05,1.14,0.68,2.15,2.24], 6.65)
    doc.add_heading("Three distinctions that prevent mistakes", level=2)
    add_bullets(doc, [
        "Venue is the goal-log value; Venue_match is the joined match-log copy used as a consistency check.",
        "Opponent_Elo is the chosen value; Elo_Source records whether it was direct or a fallback.",
        "Weighted_Goal is the base K=1 index value. The selected-K value is computed interactively and is not written back to the master CSV.",
    ])
    add_callout(doc, "Coverage note", "The master table has 29 columns. The dashboard adds six fields in memory, shown on the next page; they do not alter the source CSV.", PALE_GOLD, GOLD)

    # 22 - derived + matches 1
    new_page(doc, 22, "Field appendices B-C", "Six dashboard-derived goal fields and valid-match fields 1–10")
    derived_rows = [(group,) + field for group, fields in DERIVED_GROUPS for field in fields]
    add_table(doc, ["Group", "Field", "Type", "What it means", "Why we need it"], derived_rows, [1.05,1.14,0.68,2.15,2.24], 6.65)
    doc.add_heading("Valid-match table: fields 1–10 of 42", level=2)
    match_rows = [(group,) + field for group, fields in MATCH_GROUPS for field in fields]
    add_table(doc, ["Group", "Field", "Type", "What it means", "Why we need it"], match_rows[:10], [1.05,1.14,0.68,2.15,2.24], 6.5)

    # 23 - matches 11-26
    new_page(doc, 23, "Field appendix C", "Valid-match table: fields 11–26 of 42", "This is the analysis population; zero-goal appearances remain.")
    add_table(doc, ["Group", "Field", "Type", "What it means", "Why we need it"], match_rows[10:26], [1.05,1.14,0.68,2.15,2.24], 6.5)
    add_callout(doc, "Why many fields are not inputs", "The valid-match record keeps source context for auditing. Only Opponent_Elo, Venue, Competition_Stage, and Is_Away enter FAMD.")

    # 24 - matches 27-42
    new_page(doc, 24, "Field appendix C", "Valid-match table: fields 27–42 of 42")
    add_table(doc, ["Group", "Field", "Type", "What it means", "Why we need it"], match_rows[26:], [1.05,1.14,0.68,2.15,2.24], 6.5)
    add_callout(doc, "Denominator reminder", "Minutes is numeric and valid for all 2,201 retained rows. Gls may be zero; that is expected and required.", PALE_GOLD, GOLD)

    # 25 - glossary
    new_page(doc, 25, "Reference", "Plain-English glossary")
    glossary = [
        ("Appearance", "A match in which the player recorded positive playing time."),
        ("Bootstrap", "Resampling whole appearances with replacement to describe uncertainty in the weighted-rate gap."),
        ("Cohen's d", "A standardized difference between the two sets of appearance-level weighted rates."),
        ("Directional probability", "The share of bootstrap gaps above zero; not a p-value."),
        ("LOESS", "A descriptive local smoother used to summarize the Elo scatter; not an inferential model."),
        ("Decile", "One of ten ranked groups; D1 is easiest and D10 hardest here."),
        ("Denominator", "The quantity divided by—in this guide, all valid player minutes."),
        ("Density", "A smoothed, normalized picture of where scores are concentrated."),
        ("Dim 1", "The first FAMD axis: the single summary used as difficulty index."),
        ("Elo", "A rating that changes over time and estimates team strength."),
        ("FAMD", "Factor Analysis of Mixed Data, for continuous and categorical inputs together."),
        ("Index", "A constructed scale for comparison; not a literal count or physical unit."),
        ("K", "The exponent controlling sensitivity to the size of difficulty scores."),
        ("Normalized", "Rescaled so comparisons concern shape rather than sample size."),
        ("Per 90", "A rate scaled to 90 minutes using all valid minutes."),
        ("Sensitivity check", "A test of whether a conclusion changes when an assumption changes."),
        ("Signed power", "sign(x) × |x|ᴷ, which preserves negative-score direction."),
        ("Standardized", "Centred at approximately zero and scaled to standard-deviation units."),
        ("Weighted goal", "Here, a goal's index contribution—not a literal fractional goal."),
        ("xG", "Expected goals: a shot-quality estimate; attached but not used in this weight."),
    ]
    add_table(doc, ["Term", "Meaning in this project"], glossary, [1.65,5.45], 7.8)
    add_callout(doc, "Language choice", "Say “weighted index units per 90,” especially when values are negative. That phrasing avoids turning an abstract index into imaginary goals.")

    # 26 - speaking notes
    new_page(doc, 26, "Handoff", "How to explain this to a football fan")
    doc.add_heading("Five speaking notes", level=2)
    add_bullets(doc, [
        "Start with playing time: both raw and weighted rates use every valid minute, including scoreless matches.",
        "Separate literal facts from index results: goals and raw/90 are counts/rates; weighted/90 is an abstract context index.",
        "Explain the source of difficulty: one global FAMD combines Elo, venue, stage, and away status for both players.",
        "Explain Head-to-Head scope: competition and venue filter both goal contributions and all valid appearance minutes; penalty exclusion changes only goal contributions.",
        "End with the limits: the bootstrap quantifies conditional match-resampling uncertainty but does not repair the index's conceptual limits or declare a winner.",
    ])
    doc.add_heading("Common misunderstandings and answers", level=2)
    misunderstandings = [
        ("“A negative value means negative goals.”", "No. It is a centred index contribution; historical goal counts never change."),
        ("“The model says Elo barely matters, so opponent strength is irrelevant.”", "No. It says Elo contributes little to this particular Dim 1; other dimensions and models may represent it differently."),
        ("“Penalty exclusion should remove those match minutes.”", "No. The toggle changes goal numerators only; played minutes remain valid exposure."),
        ("“No lead change means statistical proof.”", "No. It is an exponent sensitivity result; Wave 4 separately reports the match-bootstrap distribution."),
        ("“The 96% directional probability is a p-value.”", "No. It is the fraction of resampled gaps above zero under the frozen assumptions."),
        ("“The dashboard has xG, so xG creates the weight.”", "No. xG is attached for 483 goals but is not an FAMD input or weighting term."),
        ("“The LOESS curve proves an Elo effect.”", "No. It is an unbinned descriptive smoother without a confidence band or causal interpretation."),
        ("“N/A and 0.0000 mean the same thing.”", "No. N/A means no appearances exist; zero means appearances exist but no eligible contribution remains."),
    ]
    add_table(doc, ["Misunderstanding", "Answer"], misunderstandings, [2.65,4.45], 7.5)
    add_callout(doc, "Wave 5 bridge", "The next three pages document the completed Methodology, Summary, and Raw Data tabs. They preserve the same fixed formulas, all-valid-minute denominators, and source-data contract.", PALE_TEAL, TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("USE THE EVIDENCE • STATE THE LIMITS • DO NOT DECLARE A FINAL WINNER")
    r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(TEAL)

    # 27 - Wave 5 methodology
    new_page(doc, 27, "Wave 5 methodology", "The complete analysis map, from source rows to uncertainty")
    add_table(doc, ["Step", "Dashboard evidence", "Statistical contract"], [
        ("1. Collect and clean", f"{n(f,'goal_rows'):,} goals; {n(f,'valid_match_rows'):,} valid appearances", "Keep scoreless appearances and all positive-minute exposure"),
        ("2. Build context", "One global FAMD for both players", "Inputs only: Opponent Elo, Venue, Competition Stage, Is Away"),
        ("3. Weight goals", "Stored Dim 1 score; selected signed-power K", "Missing difficulty contributes zero; negative direction is preserved"),
        ("4. Scale per 90", f"{n(f,'total_minutes'):,} total valid minutes", "Ratio of sums, never goal-scoring matches only"),
        ("5. Quantify uncertainty", f"{n(f,'wave4_reps'):,} match resamples; seed {n(f,'wave4_seed')}", "Independent within-player resampling; report the full gap distribution"),
    ], [1.25, 2.6, 3.25], 7.5)
    doc.add_heading("What creates the context index", level=2)
    add_table(doc, ["FAMD input", "Dim 1 contribution"], [
        ("Opponent Elo", f"{x(f,'famd_contrib_opponent_elo'):.4f}%"),
        ("Venue", f"{x(f,'famd_contrib_venue'):.4f}%"),
        ("Competition Stage", f"{x(f,'famd_contrib_competition_stage'):.4f}%"),
        ("Is Away", f"{x(f,'famd_contrib_is_away'):.4f}%"),
    ], [3.5, 3.6], 8.2)
    add_formula(doc, "WGᵢ(K) = sign(Dᵢ) × |Dᵢ|ᴷ", "The same signed-power rule appears in the sidebar and Methodology tab. K changes sensitivity; it does not re-run FAMD.")
    add_callout(doc, "Interpretation limit", "Dim 1 explains 25.2968% of the mixed-data variation and is a centred analytical index. Negative values are not negative goals, and a larger contribution share is not causal importance.", PALE_GOLD, GOLD)

    # 28 - Wave 5 summary
    new_page(doc, 28, "Wave 5 summary", "The fixed career reference and venue comparison")
    add_table(doc, ["Scope", "Messi apps / min", "Ronaldo apps / min", "Weighted / 90 (M / R / gap)", "Raw / 90 (M / R / gap)"], [
        ("Overall", f"{n(f,'summary_overall_messi_apps'):,} / {n(f,'summary_overall_messi_minutes'):,}", f"{n(f,'summary_overall_ronaldo_apps'):,} / {n(f,'summary_overall_ronaldo_minutes'):,}", f"{x(f,'summary_overall_messi_weighted90'):.4f} / {x(f,'summary_overall_ronaldo_weighted90'):.4f} / {x(f,'summary_overall_weighted_gap'):+.4f}", f"{x(f,'summary_overall_messi_raw90'):.4f} / {x(f,'summary_overall_ronaldo_raw90'):.4f} / {x(f,'summary_overall_raw_gap'):+.4f}"),
        ("Home", f"{n(f,'summary_venue_home_messi_apps'):,} / {n(f,'summary_venue_home_messi_minutes'):,}", f"{n(f,'summary_venue_home_ronaldo_apps'):,} / {n(f,'summary_venue_home_ronaldo_minutes'):,}", f"{x(f,'summary_venue_home_messi_weighted90'):.4f} / {x(f,'summary_venue_home_ronaldo_weighted90'):.4f} / {x(f,'summary_venue_home_weighted_gap'):+.4f}", f"{x(f,'summary_venue_home_messi_raw90'):.4f} / {x(f,'summary_venue_home_ronaldo_raw90'):.4f} / {x(f,'summary_venue_home_raw_gap'):+.4f}"),
        ("Away", f"{n(f,'summary_venue_away_messi_apps'):,} / {n(f,'summary_venue_away_messi_minutes'):,}", f"{n(f,'summary_venue_away_ronaldo_apps'):,} / {n(f,'summary_venue_away_ronaldo_minutes'):,}", f"{x(f,'summary_venue_away_messi_weighted90'):.4f} / {x(f,'summary_venue_away_ronaldo_weighted90'):.4f} / {x(f,'summary_venue_away_weighted_gap'):+.4f}", f"{x(f,'summary_venue_away_messi_raw90'):.4f} / {x(f,'summary_venue_away_ronaldo_raw90'):.4f} / {x(f,'summary_venue_away_raw_gap'):+.4f}"),
        ("Neutral", f"{n(f,'summary_venue_neutral_messi_apps'):,} / {n(f,'summary_venue_neutral_messi_minutes'):,}", f"{n(f,'summary_venue_neutral_ronaldo_apps'):,} / {n(f,'summary_venue_neutral_ronaldo_minutes'):,}", f"{x(f,'summary_venue_neutral_messi_weighted90'):.4f} / {x(f,'summary_venue_neutral_ronaldo_weighted90'):.4f} / {x(f,'summary_venue_neutral_weighted_gap'):+.4f}", f"{x(f,'summary_venue_neutral_messi_raw90'):.4f} / {x(f,'summary_venue_neutral_ronaldo_raw90'):.4f} / {x(f,'summary_venue_neutral_raw_gap'):+.4f}"),
    ], [0.8, 1.3, 1.3, 1.9, 1.8], 6.8)
    add_bullets(doc, [
        f"The fixed reference contains {n(f,'wave5_summary_rows')} ordered rows: overall, five career eras, five competition families, and three venues.",
        "The fixed table never changes. The live table updates immediately with K, penalty, competition, and venue controls; it does not wait for Update analysis.",
        "Weighted and raw columns remain side by side. N/A means missing player exposure; a real zero is printed as 0.0000.",
        "Apps and minutes reconcile to the same valid-match denominators used everywhere else.",
    ])
    add_callout(doc, "Reading rule", "Every gap is Messi minus Ronaldo. These Summary rows are descriptive; uncertainty remains on the button-frozen Inference tab.", PALE_TEAL, TEAL)

    # 29 - Wave 5 raw data and review gate
    new_page(doc, 29, "Wave 5 raw data", "Exact source transparency and the review gate")
    add_table(doc, ["Raw-data contract", "Verified value"], [
        ("Rows × columns", f"{n(f,'raw_csv_rows'):,} × {n(f,'raw_csv_columns')}"),
        ("File size", f"{n(f,'raw_csv_bytes'):,} bytes"),
        ("Source/app MD5", f[f"raw_csv_md5"]),
        ("Download behavior", "Byte-for-byte copy of goals_master_final.csv"),
        ("Dashboard controls", "No effect on rows, search, sorting, filters, or download"),
    ], [2.15, 4.95], 8.0)
    doc.add_heading("How to audit the table", level=2)
    add_bullets(doc, [
        "Use the global search to find a value across all 29 fields.",
        "Use the per-column search boxes to narrow a specific source field.",
        "Sort any header and paginate through all 1,738 goal-event rows.",
        "Scroll horizontally inside the table on narrow screens; the page itself remains contained.",
        "Download the same preserved CSV used by the tab, with no current sidebar filters applied.",
    ])
    add_callout(doc, "Stopping point", "Wave 5 is complete and verified. Methodology, Summary, and Raw Data are now fully implemented. Polish, containerization, and deployment remain in later review-gated waves and have not been started.", PALE_RED, RONALDO)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STOP AT THE WAVE 5 REVIEW GATE")
    r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(TEAL)

    normalize_output_dashes(doc)
    return doc


def verify_docx(path: Path) -> None:
    """Re-open the finished package and check required content and structures."""
    doc = Document(path)
    text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    required = [
        "The 60-second explanation", "What was collected", "The verified baseline snapshot",
        "Raw/90 = 90 × goals ÷ total valid minutes", "Dᵢ = standardized FAMD Dim 1 score",
        "WGᵢ = Dᵢ", "WGᵢ(K) = sign(Dᵢ) × |Dᵢ|ᴷ",
        "Weighted/90 = 90 × Σ WGᵢ(K) ÷ Σ minutes",
        "Gap(K) = Messi Weighted/90 - Ronaldo Weighted/90",
        "No lead change occurs", "Caveats deserve the same prominence",
        "How the career trajectories are aligned", "Opponent Elo scatter and descriptive LOESS",
        "Penalty dependency, filters, and sparse-scope rules",
        "Cumulative at t = Σ selected-K contributions through appearance t",
        "Rolling-30/90 = 90 × Σ contributions in latest 30 appearances",
        "Master goal table: fields 1-15 of 29", "Six dashboard-derived goal fields",
        "Valid-match table: fields 27-42 of 42", "Plain-English glossary",
        "How the match-level bootstrap measures uncertainty",
        "Rate* = 90 × Σ resampled match contributions ÷ Σ resampled match minutes",
        "P(Gap* > 0) = count(Gap* > 0) ÷ 10,000",
        "Career-era subgroup results", "Competition-family results and Cohen's d",
        "The complete analysis map, from source rows to uncertainty",
        "The fixed career reference and venue comparison",
        "Exact source transparency and the review gate",
        "1,738 × 29", "c43c3f995b1f301b4328c846eab2cf27",
        "0.0953", "-0.0120", "0.1999", "0.960", "0.090",
        "STOPPING POINT", "later review-gated waves", "have not been started",
        "Opponent_Elo", "Difficulty_Score", "Club_Goal", "Match Report",
    ]
    for phrase in required:
        assert phrase in text, f"Required document text missing: {phrase}"
    for token in ("TODO", "TBD", "LOREM IPSUM", "PLACEHOLDER"):
        assert token not in text.upper(), f"Placeholder token found: {token}"
    assert len(doc.tables) >= 25
    with zipfile.ZipFile(path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        footer_xml = "".join(
            zf.read(name).decode("utf-8") for name in zf.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )
    assert "w:tblHeader" in document_xml, "Repeated table headers are missing"
    assert " PAGE " in footer_xml, "Page-number field is missing"
    assert document_xml.count('w:val="Heading1"') >= 28, "Heading hierarchy is incomplete"


def main() -> None:
    facts = extract_facts()
    verify_facts(facts)
    doc = build_document(facts)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    verify_docx(OUTPUT)
    print(f"Created and verified: {OUTPUT}")
    print("Facts: 1,738 goals; 2,201 appearances; 181,081 minutes; 1,063 scoreless appearances")
    print("Scope: verified through Wave 5; Wave 6 and later work explicitly not completed")


if __name__ == "__main__":
    main()
